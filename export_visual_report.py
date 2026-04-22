from __future__ import annotations

import argparse
import base64
import subprocess
import sys
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


COMMON_BROWSER_PATHS = (
    Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
    Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
    Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
    Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
)

OUTPUT_SELECTORS = (
    ".jp-RenderedHTMLCommon",
    ".jp-RenderedImage",
    ".jp-OutputArea-output pre",
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export a notebook HTML report to DOCX and/or PDF."
    )
    parser.add_argument("input_html", type=Path, help="Path to the HTML report.")
    parser.add_argument(
        "--docx",
        type=Path,
        help="Output DOCX path. Defaults to the HTML path with .docx extension.",
    )
    parser.add_argument(
        "--pdf",
        type=Path,
        help="Output PDF path. Defaults to the HTML path with .pdf extension.",
    )
    parser.add_argument(
        "--max-image-width",
        type=float,
        default=6.5,
        help="Maximum image width in inches for the DOCX export.",
    )
    parser.add_argument(
        "--no-docx",
        action="store_true",
        help="Skip DOCX export.",
    )
    parser.add_argument(
        "--no-pdf",
        action="store_true",
        help="Skip PDF export.",
    )
    return parser.parse_args()


def resolve_output(input_html: Path, output: Path | None, extension: str) -> Path:
    return output if output is not None else input_html.with_suffix(extension)


def find_browser() -> Path | None:
    for candidate in COMMON_BROWSER_PATHS:
        if candidate.exists():
            return candidate
    return None


def load_soup(path: Path) -> BeautifulSoup:
    return BeautifulSoup(path.read_text(encoding="utf-8"), "lxml")


def iter_output_blocks(soup: BeautifulSoup) -> Iterable[Tag]:
    seen_ids: set[int] = set()
    for selector in OUTPUT_SELECTORS:
        for node in soup.select(selector):
            marker = id(node)
            if marker not in seen_ids:
                seen_ids.add(marker)
                yield node


def clean_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    cleaned = clean_text(text)
    if cleaned:
        document.add_paragraph(cleaned, style=style)


def heading_level(tag_name: str) -> int:
    if tag_name.startswith("h") and len(tag_name) == 2 and tag_name[1].isdigit():
        return min(max(int(tag_name[1]), 1), 4)
    return 1


def add_table(document: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return

    max_cols = max(
        sum(int(cell.get("colspan", 1)) for cell in row.find_all(["th", "td"]))
        for row in rows
    )
    table = document.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"

    for row_tag in rows:
        row = table.add_row().cells
        col_index = 0
        for cell_tag in row_tag.find_all(["th", "td"]):
            text = clean_text(cell_tag.get_text(" ", strip=True))
            cell = row[col_index]
            cell.text = text
            col_index += int(cell_tag.get("colspan", 1))


def add_image(document: Document, img_tag: Tag, max_width_inches: float) -> None:
    src = img_tag.get("src", "")
    if not src.startswith("data:image/"):
        return

    try:
        _, encoded = src.split(",", 1)
        image_bytes = base64.b64decode(encoded)
    except Exception:
        return

    picture = BytesIO(image_bytes)
    document.add_picture(picture, width=Inches(max_width_inches))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    alt_text = img_tag.get("alt", "").strip()
    if alt_text and alt_text != "No description has been provided for this image":
        caption = document.add_paragraph(alt_text)
        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def render_html_fragment(document: Document, node: Tag, max_width_inches: float) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = clean_text(str(child))
            if text:
                document.add_paragraph(text)
            continue

        if not isinstance(child, Tag):
            continue

        if child.name in {"h1", "h2", "h3", "h4"}:
            document.add_heading(clean_text(child.get_text(" ", strip=True)), level=heading_level(child.name))
        elif child.name == "p":
            add_paragraph(document, child.get_text(" ", strip=True))
        elif child.name in {"ul", "ol"}:
            style = "List Bullet" if child.name == "ul" else "List Number"
            for item in child.find_all("li", recursive=False):
                add_paragraph(document, item.get_text(" ", strip=True), style=style)
        elif child.name == "table":
            add_table(document, child)
            document.add_paragraph("")
        elif child.name == "img":
            add_image(document, child, max_width_inches)
        elif child.name == "pre":
            text = child.get_text("\n", strip=True)
            if text:
                document.add_paragraph(text)
        else:
            render_html_fragment(document, child, max_width_inches)


def export_docx(input_html: Path, output_docx: Path, max_width_inches: float) -> None:
    soup = load_soup(input_html)
    document = Document()
    title = soup.title.string.strip() if soup.title and soup.title.string else input_html.stem
    document.add_heading(title, level=0)

    blocks = list(iter_output_blocks(soup))
    if not blocks:
        raise RuntimeError(
            "No notebook output blocks were found in the HTML. "
            "Export an executed notebook HTML first."
        )

    for block in blocks:
        render_html_fragment(document, block, max_width_inches)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def export_pdf(input_html: Path, output_pdf: Path) -> None:
    browser = find_browser()
    if browser is None:
        raise RuntimeError(
            "No supported Chromium browser was found. Install Chrome or Edge to export PDF."
        )

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    file_url = input_html.resolve().as_uri()
    command = [
        str(browser),
        "--headless",
        "--disable-gpu",
        "--no-sandbox",
        f"--print-to-pdf={output_pdf.resolve()}",
        file_url,
    ]
    subprocess.run(command, check=True)


def main() -> int:
    args = parse_args()
    input_html = args.input_html.resolve()
    docx_path = resolve_output(input_html, args.docx, ".docx")
    pdf_path = resolve_output(input_html, args.pdf, ".pdf")
    export_docx_enabled = not args.no_docx
    export_pdf_enabled = not args.no_pdf

    if not input_html.exists():
        print(f"Input HTML not found: {input_html}", file=sys.stderr)
        return 1
    if not export_docx_enabled and not export_pdf_enabled:
        print("Nothing to export. Remove --no-docx or --no-pdf.", file=sys.stderr)
        return 1

    try:
        if export_docx_enabled:
            export_docx(input_html, docx_path, args.max_image_width)
        if export_pdf_enabled:
            export_pdf(input_html, pdf_path)
    except subprocess.CalledProcessError as exc:
        print(f"PDF export failed: {exc}", file=sys.stderr)
        return exc.returncode or 1
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        return 1

    if export_docx_enabled:
        print(f"DOCX saved to {docx_path}")
    if export_pdf_enabled:
        print(f"PDF saved to {pdf_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
